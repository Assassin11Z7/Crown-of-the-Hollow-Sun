#!/usr/bin/env python3
"""Assemble markdown chapters into JSON (for the reader) and .docx (for Word)."""

from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path("/workspace")
CONTENT = ROOT / "content"
PUBLIC = ROOT / "public"
BOOKS_DIR = PUBLIC / "books"
JSON_DIR = PUBLIC / "content"

INK = RGBColor(0x2A, 0x24, 0x1C)
MUTED = RGBColor(0x6A, 0x5E, 0x4E)
BRASS = RGBColor(0x8A, 0x62, 0x30)

FRONT = {
    "ignition": {
        "id": "ignition",
        "volume": 1,
        "title": "Ignition",
        "series": "Crown of the Hollow Sun",
        "subtitle": "The Boy Who Brought Noon",
        "dedication": "For the cold ones who still shared their last coal.",
        "epigraph": {
            "text": "The sun did not die. It learned to hide inside a boy.",
            "attribution": "Nightlands proverb, unrecorded",
        },
        "docx_name": "Crown-of-the-Hollow-Sun-1-Ignition.docx",
        "folder": "book1",
    },
    "ascension": {
        "id": "ascension",
        "volume": 2,
        "title": "Ascension",
        "series": "Crown of the Hollow Sun",
        "subtitle": "The Dream Becoming Lucid",
        "dedication": "For the ones who defected and then had to live with it.",
        "epigraph": {
            "text": "He is not the first fire. He is the first fire that asked permission.",
            "attribution": "From the Nightlands memory-ice, unnamed",
        },
        "docx_name": "Crown-of-the-Hollow-Sun-2-Ascension.docx",
        "folder": "book2",
    },
    "eclipse": {
        "id": "eclipse",
        "volume": 3,
        "title": "Eclipse",
        "series": "Crown of the Hollow Sun",
        "subtitle": "A Billion Small Flames",
        "dedication": "For anyone who was told there wasn’t enough light to go around.",
        "epigraph": {
            "text": "Salvation that requires a famine is not salvation. It is administration.",
            "attribution": "Ren Kael, later, to no one in particular",
        },
        "docx_name": "Crown-of-the-Hollow-Sun-3-Eclipse.docx",
        "folder": "book3",
    },
}


def set_run_font(run, name="Garamond", size=11, italic=False, bold=False, color=INK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    run.font.color.rgb = color


def add_page_number(paragraph):
    run1 = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run1._r.append(fld1)
    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)
    run3 = paragraph.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld2)


def parse_chapter(path: Path) -> dict:
    raw = path.read_text(encoding="utf-8")
    lines = raw.replace("\r\n", "\n").split("\n")
    number = None
    title = None
    body_start = 0
    for i, line in enumerate(lines):
        m = re.match(r"^#\s+Chapter\s+(\d+)\s*$", line.strip(), re.I)
        if m:
            number = int(m.group(1))
            body_start = i + 1
            continue
        m2 = re.match(r"^##\s+(.+?)\s*$", line.strip())
        if m2 and title is None:
            title = m2.group(1).strip()
            body_start = i + 1
            continue
    if number is None:
        m = re.search(r"(\d+)", path.stem)
        number = int(m.group(1)) if m else 0
    if title is None:
        title = path.stem
    paras: list[str] = []
    buf: list[str] = []
    for line in lines[body_start:]:
        s = line.rstrip()
        if s.strip() in {"*", "* * *", "***", "⁂"}:
            if buf:
                paras.append(" ".join(x.strip() for x in buf if x.strip()))
                buf = []
            paras.append("⁂")
            continue
        if s.strip() == "":
            if buf:
                paras.append(" ".join(x.strip() for x in buf if x.strip()))
                buf = []
            continue
        if s.startswith("#"):
            continue
        buf.append(s)
    if buf:
        paras.append(" ".join(x.strip() for x in buf if x.strip()))
    paras = [p for p in paras if p]
    return {"number": number, "title": title, "paragraphs": paras}


def word_count(chapters: list[dict]) -> int:
    n = 0
    for ch in chapters:
        for p in ch["paragraphs"]:
            n += len(re.findall(r"\S+", p))
    return n


def add_mixed_runs(paragraph, text: str, size=11):
    parts = re.split(r"(\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        italic = part.startswith("*") and part.endswith("*") and len(part) > 2
        content = part[1:-1] if italic else part
        run = paragraph.add_run(content)
        set_run_font(run, size=size, italic=italic, color=INK)


def build_docx(meta: dict, chapters: list[dict], dest: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(6)
        section.page_height = Inches(9)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(p)
        if p.runs:
            set_run_font(p.runs[0], size=9, color=MUTED)
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(f"{meta['series']}  ·  {meta['title']}")
        set_run_font(run, size=9, italic=True, color=MUTED)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta["series"].upper())
    set_run_font(r, size=12, color=BRASS)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta["title"])
    set_run_font(r, size=36, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Book {['I', 'II', 'III'][meta['volume'] - 1]}")
    set_run_font(r, size=14, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta["subtitle"])
    set_run_font(r, size=14, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("A novel")
    set_run_font(r, size=12, color=MUTED)

    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"{meta['series']}: {meta['title']}")
    set_run_font(r, size=11, italic=True)
    p = doc.add_paragraph()
    r = p.add_run("This is a work of fiction.")
    set_run_font(r, size=10, color=MUTED)
    p = doc.add_paragraph()
    r = p.add_run("Set in Garamond for reading in Microsoft Word, Pages, and Google Docs.")
    set_run_font(r, size=10, color=MUTED)

    doc.add_page_break()
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta["dedication"])
    set_run_font(r, size=13, italic=True)

    doc.add_page_break()
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"“{meta['epigraph']['text']}”")
    set_run_font(r, size=13, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— " + meta["epigraph"]["attribution"])
    set_run_font(r, size=10, color=MUTED)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Contents")
    set_run_font(r, size=20)
    p.paragraph_format.space_after = Pt(18)
    for ch in chapters:
        p = doc.add_paragraph()
        r = p.add_run(f"Chapter {ch['number']}")
        set_run_font(r, size=10, color=MUTED)
        p = doc.add_paragraph()
        r = p.add_run(ch["title"])
        set_run_font(r, size=13)
        p.paragraph_format.space_after = Pt(10)

    doc.add_page_break()
    for ci, ch in enumerate(chapters):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"CHAPTER {ch['number']}")
        set_run_font(r, size=11, color=BRASS)
        p.paragraph_format.space_after = Pt(6)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(ch["title"])
        set_run_font(r, size=22)
        p.paragraph_format.space_after = Pt(22)
        first = True
        for para in ch["paragraphs"]:
            if para.strip() in {"⁂", "*", "* * *"}:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("⁂")
                set_run_font(r, size=12, color=BRASS)
                first = True
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.first_line_indent = Inches(0 if first else 0.28)
            first = False
            add_mixed_runs(p, para, size=11)
        if ci != len(chapters) - 1:
            doc.add_page_break()

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))


def assemble(book_id: str) -> dict | None:
    meta = FRONT[book_id]
    folder = CONTENT / meta["folder"]
    files = sorted(folder.glob("*.md"))
    chapters = [parse_chapter(p) for p in files]
    chapters = [c for c in chapters if c["paragraphs"]]
    chapters.sort(key=lambda c: c["number"])
    if not chapters:
        print(f"{book_id}: no chapters yet")
        return None
    payload = {
        "id": meta["id"],
        "volume": meta["volume"],
        "title": meta["title"],
        "series": meta["series"],
        "subtitle": meta["subtitle"],
        "dedication": meta["dedication"],
        "epigraph": meta["epigraph"],
        "chapters": chapters,
        "wordCount": word_count(chapters),
    }
    JSON_DIR.mkdir(parents=True, exist_ok=True)
    (JSON_DIR / f"{book_id}.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    build_docx(meta, chapters, BOOKS_DIR / meta["docx_name"])
    print(f"{book_id}: {len(chapters)} chapters, {payload['wordCount']} words")
    return payload


def main():
    BOOKS_DIR.mkdir(parents=True, exist_ok=True)
    JSON_DIR.mkdir(parents=True, exist_ok=True)
    built = []
    for book_id in ("ignition", "ascension", "eclipse"):
        payload = assemble(book_id)
        if payload:
            built.append(FRONT[book_id]["docx_name"])
    if len(built) == 3:
        zip_path = BOOKS_DIR / "Crown-of-the-Hollow-Sun-Trilogy.zip"
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for name in built:
                zf.write(BOOKS_DIR / name, name)
        print(f"zip: {zip_path}")


if __name__ == "__main__":
    main()
